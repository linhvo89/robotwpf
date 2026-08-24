from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_khoi_dong_nguon.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_ap_suat_khi.docx"
image = Path(r"C:\Users\LinhVo\OneDrive\Desktop\anh\apsuatkhi.PNG")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "3.2 Kiểm tra điều kiện sẵn sàng"
    and paragraph.style.name == "Heading 2"
)
anchor.text = "3.3 Kiểm tra điều kiện sẵn sàng"
anchor.style = "Heading 2"

created = []

heading = doc.add_paragraph("3.2 Kiểm tra và cài đặt áp suất khí", style="Heading 2")
created.append(heading._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(3.0))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 3. Màn hình hiển thị áp suất khí của hệ thống")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

instructions = [
    ("1. Kiểm tra giá trị: ", "Đọc áp suất trên màn hình hiển thị của bộ điều áp."),
    ("2. Cài đặt áp suất: ", "Điều chỉnh bộ lọc–điều áp để áp suất nằm trong khoảng từ 0,38 MPa đến 0,60 MPa."),
    ("3. Xác nhận ổn định: ", "Chờ giá trị ổn định trong dải cho phép trước khi khởi động chu trình Auto."),
]
for label, detail in instructions:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(8)
run = note.add_run("LƯU Ý: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
note.add_run(
    "Không vận hành nếu áp suất thấp hơn 0,38 MPa hoặc cao hơn 0,60 MPa. "
    "Nếu không điều chỉnh được hoặc áp suất dao động bất thường, dừng hệ thống và báo kỹ thuật."
)
created.append(note._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

for paragraph in doc.paragraphs:
    if "Áp suất khí đạt yêu cầu nội bộ" in paragraph.text:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = (
            "4. Kiểm tra khí nén: Áp suất khí ổn định trong khoảng 0,38–0,60 MPa "
            "và không có rò rỉ bất thường."
        )

doc.core_properties.comments = "Bổ sung dải áp suất khí 0,38–0,60 MPa ngày 11/08/2026"
doc.save(output)
print(output)
