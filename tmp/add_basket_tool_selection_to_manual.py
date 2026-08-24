from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_cai_dat_toc_do_robot.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_chon_basket_va_tool.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-2d6bbdbc-8954-4ebe-a64e-798776ee6499.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("7.5 Chọn Basket và đầu Tool chạy", style="Heading 2")
created.append(heading._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(6.3))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 7. Chọn Basket và các đầu Tool tham gia chu trình")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

for label, detail in [
    ("Basket1", "Chỉ chạy Basket 1; hệ thống sử dụng trạng thái Basket 1 và Camera 1."),
    ("Basket2", "Chỉ chạy Basket 2; hệ thống sử dụng trạng thái Basket 2 và Camera 2."),
    ("Both", "Chạy cả hai Basket; hệ thống xử lý Basket 1 trước, sau đó chuyển sang Basket 2."),
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

tool_text = doc.add_paragraph()
tool_text.paragraph_format.space_after = Pt(6)
run = tool_text.add_run("CHỌN ĐẦU TOOL: ")
run.bold = True
run.font.color.rgb = RGBColor(22, 74, 138)
tool_text.add_run(
    "Đánh dấu Tool 1, Tool 2 và/hoặc Tool 3 muốn sử dụng. Có thể chọn một Tool, nhiều Tool hoặc cả ba Tool. "
    "Hệ thống chỉ dùng các đầu Tool đang được đánh dấu."
)
created.append(tool_text._p)

steps = [
    ("1. Chọn Basket: ", "Mở danh sách Select Basket và chọn Basket1, Basket2 hoặc Both theo kế hoạch sản xuất."),
    ("2. Chọn Tool: ", "Đánh dấu các đầu Tool cần chạy; phải chọn ít nhất một Tool."),
    ("3. Kiểm tra sẵn sàng: ", "Xác nhận Basket đã chọn có khay/sản phẩm, cảm biến sẵn sàng và H1–H3 của các Tool phù hợp."),
    ("4. Chạy thử: ", "Ở lần chạy đầu sau khi đổi lựa chọn, đặt tốc độ thấp và quan sát chu trình gắp/thả."),
]
for label, detail in steps:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("LƯU Ý: ")
run.bold = True
run.font.color.rgb = RGBColor(176, 122, 0)
warning.add_run(
    "Nếu chọn Both, cả Basket 1 và Basket 2 phải ở trạng thái sẵn sàng. "
    "Nếu bỏ chọn toàn bộ Tool, robot sẽ không thực hiện thao tác gắp."
)
created.append(warning._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung hướng dẫn chọn Basket và Tool ngày 11/08/2026"
doc.save(output)
print(output)
