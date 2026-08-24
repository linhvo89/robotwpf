from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_ap_suat_khi.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_nut_emergency.docx"
image = Path(r"C:\Users\LinhVo\OneDrive\Desktop\anh\Emergence.PNG")

doc = Document(source)

# Đánh lại số hình theo thứ tự xuất hiện sau khi thêm ảnh Emergency Stop.
caption_updates = {
    "Hình 3. Màn hình hiển thị áp suất khí của hệ thống": "Hình 4. Màn hình hiển thị áp suất khí của hệ thống",
    "Hình 2. Aptomat nguồn và đèn BÁO NGUỒN trên tủ điều khiển": "Hình 3. Aptomat nguồn và đèn BÁO NGUỒN trên tủ điều khiển",
    "Hình 1. Nhóm nút vận hành chính trên màn hình Home": "Hình 2. Nhóm nút vận hành chính trên màn hình Home",
}
for paragraph in doc.paragraphs:
    if paragraph.text.strip() in caption_updates:
        replacement = caption_updates[paragraph.text.strip()]
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = replacement

anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "1.2 Khi nào phải nhấn Emergency Stop"
    and paragraph.style.name == "Heading 2"
)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(1.45))
picture.paragraph_format.space_after = Pt(3)

caption = doc.add_paragraph("Hình 1. Nút Emergency Stop của hệ thống")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(6)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)

current = anchor._p
for element in (picture._p, caption._p):
    current.addnext(element)
    current = element

doc.core_properties.comments = "Bổ sung ảnh nút Emergency Stop ngày 11/08/2026"
doc.save(output)
print(output)
