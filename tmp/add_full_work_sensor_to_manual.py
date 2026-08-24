from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_chon_basket_va_tool.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_sensor_full_work.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-181c5d93-007f-412a-808b-267353fe89ca.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("7.6 Chọn máy nhận sản phẩm và cảnh báo Full Work", style="Heading 2")
created.append(heading._p)

intro = doc.add_paragraph(
    "Trong mục Sensor Full Work, chọn máy mà robot sẽ gắp và thả sản phẩm vào. "
    "Lựa chọn này xác định cảm biến được dùng để giám sát trạng thái đầy."
)
created.append(intro._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(5.4))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 8. Chọn cảm biến Full Work của máy nhận sản phẩm")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

for label, detail in [
    ("Máy1", "Robot thả sản phẩm vào Máy 1 và giám sát cảm biến X0/20480."),
    ("Máy2", "Robot thả sản phẩm vào Máy 2 và giám sát cảm biến X1/20481."),
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

sequence = [
    ("1. Chọn máy: ", "Mở danh sách Sensor Full Work và chọn đúng Máy1 hoặc Máy2 theo vị trí robot sẽ thả sản phẩm."),
    ("2. Khi máy đầy: ", "Sau hai lần thả liên tiếp mà cảm biến máy đã chọn vẫn báo đầy, hệ thống phát trạng thái Full Work."),
    ("3. Robot chờ an toàn: ", "Robot quay về HomePose, tạm dừng chu trình và đèn xanh nhấp nháy để báo máy đầy."),
    ("4. Lấy sản phẩm ra: ", "Người vận hành lấy sản phẩm khỏi máy nhận và bảo đảm vùng máy an toàn."),
    ("5. Tiếp tục: ", "Khi cảm biến trở về trạng thái trống, đèn xanh sáng liên tục và chương trình tự động tiếp tục."),
]
for label, detail in sequence:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(8)
run = note.add_run("LƯU Ý: ")
run.bold = True
run.font.color.rgb = RGBColor(176, 122, 0)
note.add_run(
    "Phải chọn đúng máy nhận trước khi Start. Không đưa tay vào vùng robot khi chỉ mới thấy cảnh báo Full Work; "
    "thực hiện quy trình an toàn tại máy trước khi lấy sản phẩm."
)
created.append(note._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung hướng dẫn Sensor Full Work ngày 11/08/2026"
doc.save(output)
print(output)
