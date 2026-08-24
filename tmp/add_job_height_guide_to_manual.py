from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_nut_emergency.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_chon_job_H1_H2_H3.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-93caf0d9-4a3c-496f-9208-6c869eb164a4.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "4.1 Khởi động chu trình"
    and paragraph.style.name == "Heading 2"
)

created = []

heading = doc.add_paragraph("4.1.1 Chọn Job/Model và cài đặt H1–H3", style="Heading 3")
created.append(heading._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(6.3))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 5. Chọn Job/Model và các giá trị H1, H2, H3 trên màn hình Home")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

steps = [
    ("1. Hiển thị danh sách Job: ", "Nếu danh sách chưa hiển thị, nhấn nút Model."),
    ("2. Chọn Job chạy: ", "Nhấp vào dòng Job Name cần chạy, kiểm tra đúng tên Job rồi chọn Yes trong hộp xác nhận."),
    ("3. Xác nhận Job: ", "Tên Job đã chọn phải hiển thị màu xanh bên cạnh chữ “Vận hành” trước khi nhấn Start."),
]
for label, detail in steps:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

for label, detail in [
    ("H1", "Độ sâu gắp của Tool 1 / đầu hút xi lanh 1."),
    ("H2", "Độ sâu gắp của Tool 2 / đầu hút xi lanh 2."),
    ("H3", "Độ sâu gắp của Tool 3 / đầu hút xi lanh 3."),
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

effect = doc.add_paragraph()
effect.paragraph_format.space_after = Pt(6)
run = effect.add_run("QUY TẮC ĐIỀU CHỈNH: ")
run.bold = True
run.font.color.rgb = RGBColor(22, 74, 138)
effect.add_run(
    "Giá trị H càng lớn thì robot hạ đầu hút xuống càng sâu; giá trị H càng nhỏ thì robot hạ ít hơn. "
    "Chỉ cài đặt trong giới hạn từ 0 đến 20."
)
created.append(effect._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("CẢNH BÁO: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
warning.add_run(
    "Không đặt H quá lớn vì đầu hút có thể ép mạnh vào sản phẩm, đồ gá hoặc gây va chạm. "
    "Sau mỗi lần thay đổi, chạy thử ở Manual với tốc độ thấp và tăng/giảm từng bước nhỏ."
)
created.append(warning._p)

current = anchor._p
for element in created:
    current.addnext(element)
    current = element

doc.core_properties.comments = "Bổ sung hướng dẫn chọn Job và H1-H3 ngày 11/08/2026"
doc.save(output)
print(output)
