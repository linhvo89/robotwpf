from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_setsensor.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_giai_thich_nut_robot.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-af4cade3-d351-4cdd-a866-d4300ad199d4.png")

doc = Document(source)

# Chèn một hình mới trước Chương 7 nên đánh lại số các hình phía sau.
caption_updates = {
    "Hình 10. Tùy chọn SetSensor trong nhóm System": "Hình 11. Tùy chọn SetSensor trong nhóm System",
    "Hình 9. Bảng Offset điểm hút cho Basket 1, Basket 2 và Tool 1–3": "Hình 10. Bảng Offset điểm hút cho Basket 1, Basket 2 và Tool 1–3",
    "Hình 8. Chọn cảm biến Full Work của máy nhận sản phẩm": "Hình 9. Chọn cảm biến Full Work của máy nhận sản phẩm",
    "Hình 7. Chọn Basket và các đầu Tool tham gia chu trình": "Hình 8. Chọn Basket và các đầu Tool tham gia chu trình",
    "Hình 6. Nhóm cài đặt tốc độ robot trong màn hình Settings": "Hình 7. Nhóm cài đặt tốc độ robot trong màn hình Settings",
}
for paragraph in doc.paragraphs:
    if paragraph.text.strip() in caption_updates:
        replacement = caption_updates[paragraph.text.strip()]
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = replacement

anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "7. Cài đặt Job và vị trí robot"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("6.3 Các nút điều khiển và trạng thái Robot", style="Heading 2")
created.append(heading._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(6.3))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 6. Nhóm trạng thái và nút điều khiển Robot trên màn hình Manual")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

state_intro = doc.add_paragraph()
state_intro.paragraph_format.space_after = Pt(5)
run = state_intro.add_run("TRÌNH TỰ TRẠNG THÁI: ")
run.bold = True
run.font.color.rgb = RGBColor(22, 74, 138)
state_intro.add_run("Powered on → Controller initialized → Disable → Standby.")
created.append(state_intro._p)

for label, detail in [
    ("Powered on", "Nguồn điều khiển robot đã được cấp."),
    ("Controller initialized", "Bộ điều khiển đã khởi tạo và sẵn sàng quản lý robot."),
    ("Disable", "Servo đang tắt; robot chưa được phép chuyển động theo lệnh."),
    ("Standby", "Servo đã Enable; robot sẵn sàng nhận lệnh chuyển động khi các điều kiện an toàn hợp lệ."),
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

for label, detail in [
    ("POWER ON / INITIALIZE / ENABLE / DISABLE", "Nút bên trái tự đổi tên theo trạng thái. Nhấn lần lượt POWER ON để cấp nguồn, INITIALIZE để khởi tạo Controller, ENABLE để bật Servo. Khi robot đã Standby, nút đổi thành DISABLE để tắt Servo có kiểm soát."),
    ("POWER OFF", "Thực hiện quy trình tắt robot an toàn: tắt Servo, đóng Controller và ngắt nguồn điều khiển 48 V. Chỉ nhấn khi robot đã dừng và không mang sản phẩm nguy hiểm."),
    ("Reset Robot", "Tắt các ngõ ra điều khiển và gửi lệnh xóa lỗi robot. Chỉ Reset sau khi đã xác định và xử lý nguyên nhân lỗi/Emergency."),
    ("Status Robot", "Đọc trạng thái chi tiết của robot và ghi kết quả vào Machine Log để kiểm tra nguồn, Servo, chuyển động và lỗi."),
]:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("CẢNH BÁO: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
warning.add_run(
    "Không ENABLE robot khi có người hoặc vật lạ trong vùng chuyển động. Không dùng Reset Robot để bỏ qua lỗi chưa rõ nguyên nhân. "
    "Luôn kiểm tra Machine Log sau mỗi thao tác không thành công."
)
created.append(warning._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung giải thích nút Robot, không gồm Free Drive, ngày 11/08/2026"
doc.save(output)
print(output)
