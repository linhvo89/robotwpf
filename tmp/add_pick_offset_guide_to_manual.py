from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_sensor_full_work.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_offset_diem_hut.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-90331323-276e-441e-bd2a-9e10c23a9097.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("7.7 Hiệu chỉnh lệch hút bằng Offset", style="Heading 2")
created.append(heading._p)

intro = doc.add_paragraph(
    "Khi đầu hút tiếp cận lệch tâm sản phẩm, có thể dùng Delta X và Delta Y để hiệu chỉnh riêng "
    "cho từng Basket, từng Tool và từng vùng tọa độ X của sản phẩm."
)
created.append(intro._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(6.3))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 9. Bảng Offset điểm hút cho Basket 1, Basket 2 và Tool 1–3")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

for label, detail in [
    ("Delta X", "Được cộng trực tiếp vào tọa độ X của điểm hút. Số dương dịch theo chiều +X; số âm dịch theo chiều −X."),
    ("Delta Y", "Được cộng trực tiếp vào tọa độ Y của điểm hút. Số dương dịch theo chiều +Y; số âm dịch theo chiều −Y."),
    ("X < PickProductPose.X", "Dùng Offset của dòng này khi tọa độ X sản phẩm nhỏ hơn X của PickProductPose."),
    ("X ≥ PickProductPose.X", "Dùng Offset của dòng này khi tọa độ X sản phẩm lớn hơn hoặc bằng X của PickProductPose."),
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

steps = [
    ("1. Xác định trường hợp lệch: ", "Ghi nhận Basket đang chạy, Tool đang hút, vùng X của sản phẩm và hướng lệch thực tế."),
    ("2. Chọn đúng dòng: ", "Tìm đúng Basket 1/2, Tool 1/2/3 và điều kiện X < hoặc X ≥ PickProductPose.X."),
    ("3. Nhập Offset nhỏ: ", "Điều chỉnh Delta X/Delta Y từng bước nhỏ; nhấp ra ngoài ô để phần mềm nhận giá trị."),
    ("4. Lưu dữ liệu: ", "Nhấn Save Offset và kiểm tra thông báo lưu thành công."),
    ("5. Chạy xác nhận: ", "Chạy thử Manual ở tốc độ thấp, quan sát điểm hút rồi tiếp tục hiệu chỉnh nếu cần."),
]
for label, detail in steps:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("CẢNH BÁO: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
warning.add_run(
    "Không nhập Offset lớn ngay từ đầu. Nếu chỉnh sai dấu làm độ lệch tăng, đưa giá trị về trước đó và đổi chiều hiệu chỉnh. "
    "Nếu phải dùng Offset lớn hoặc nhiều vị trí đều lệch, kiểm tra lại camera, calibration, đầu hút và cơ khí thay vì tiếp tục bù Offset."
)
created.append(warning._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung hướng dẫn Offset điểm hút ngày 11/08/2026"
doc.save(output)
print(output)
