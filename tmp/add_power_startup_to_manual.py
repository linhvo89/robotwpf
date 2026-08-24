from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_anh_trang_van_hanh.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_khoi_dong_nguon.docx"
image = Path(r"C:\Users\LinhVo\OneDrive\Desktop\anh\1.PNG")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "3. Kiểm tra trước khi khởi động"
    and paragraph.style.name == "Heading 1"
)

created = []

heading = doc.add_paragraph("3.1 Cấp nguồn và khởi động hệ thống", style="Heading 2")
created.append(heading._p)

intro = doc.add_paragraph(
    "Thực hiện lần lượt các bước dưới đây khi bắt đầu ca hoặc sau khi hệ thống đã được tắt hoàn toàn."
)
created.append(intro._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(4.4))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 2. Aptomat nguồn và đèn BÁO NGUỒN trên tủ điều khiển")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

steps = [
    ("1. Kiểm tra đèn báo nguồn: ", "Xác nhận đèn BÁO NGUỒN màu xanh đang sáng."),
    ("2. Bật aptomat nguồn: ", "Gạt aptomat nguồn sang vị trí ON để cấp nguồn khởi động robot và máy tính điều khiển."),
    ("3. Mở phần mềm: ", "Sau khi máy tính khởi động, mở phần mềm KBOT trên máy tính."),
    ("4. Chờ hệ thống sẵn sàng: ", "Quá trình khởi động robot và phần mềm có thể mất khoảng 5 phút. Chờ màn hình KBOT hiển thị ổn định và không còn thông báo đang khởi tạo."),
]
for label, detail in steps:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(8)
run = note.add_run("LƯU Ý: ")
run.bold = True
run.font.color.rgb = RGBColor(176, 122, 0)
note.add_run(
    "Không nhấn Start, Home hoặc điều khiển Manual trong thời gian hệ thống đang khởi động. "
    "Chỉ tiếp tục khi robot, máy tính và phần mềm đã khởi động hoàn tất."
)
created.append(note._p)

subheading = doc.add_paragraph("3.2 Kiểm tra điều kiện sẵn sàng", style="Heading 2")
created.append(subheading._p)

current = anchor._p
for element in created:
    current.addnext(element)
    current = element

doc.core_properties.comments = "Bổ sung quy trình cấp nguồn và khởi động ngày 11/08/2026"
doc.save(output)
print(output)
