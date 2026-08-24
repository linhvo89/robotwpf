from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_chon_job_H1_H2_H3.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_cai_dat_toc_do_robot.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-18d182ea-bed9-4a4a-abc1-e2b857081175.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("7.4 Cài đặt tốc độ robot", style="Heading 2")
created.append(heading._p)

intro = doc.add_paragraph(
    "Trong màn hình Settings, chọn giá trị tại từng danh sách trong nhóm “Tốc độ robot”. "
    "Phần mềm cho phép chọn từ 0,05 đến 1,00, theo bước 0,05."
)
created.append(intro._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(6.0))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 6. Nhóm cài đặt tốc độ robot trong màn hình Settings")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

items = [
    ("Tốc độ Robot chụp ảnh", "Tốc độ di chuyển đến vị trí chụp/kiểm tra camera."),
    ("Tốc độ Robot hút", "Tốc độ di chuyển khi tiếp cận và thực hiện gắp sản phẩm."),
    ("Tốc độ Robot đi thả 1", "Tốc độ đi đến vị trí thả thứ nhất."),
    ("Tốc độ Robot đi lên vị trí thả từ 1 đến 5", "Tốc độ di chuyển giữa các vị trí thả."),
    ("Tốc độ Robot quay về sau khi thả", "Tốc độ robot quay về sau khi hoàn thành thao tác thả."),
]
for label, detail in items:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(detail)
    created.append(paragraph._p)

steps = doc.add_paragraph()
steps.paragraph_format.space_after = Pt(6)
run = steps.add_run("CÁCH CÀI ĐẶT: ")
run.bold = True
run.font.color.rgb = RGBColor(22, 74, 138)
steps.add_run(
    "Nhấp mũi tên bên phải từng dòng, chọn tốc độ phù hợp rồi chạy thử. "
    "Giá trị càng lớn thì robot chạy càng nhanh; giá trị càng nhỏ thì robot chạy càng chậm."
)
created.append(steps._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("CẢNH BÁO: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
warning.add_run(
    "Khi chạy thử hoặc sau khi thay đổi điểm/H1–H3, đặt tốc độ thấp 0,05 và tăng từng bước nhỏ. "
    "Chỉ tăng tốc sau khi đã xác nhận quỹ đạo, khoảng hở và thao tác gắp/thả an toàn."
)
created.append(warning._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung hướng dẫn cài đặt tốc độ robot ngày 11/08/2026"
doc.save(output)
print(output)
