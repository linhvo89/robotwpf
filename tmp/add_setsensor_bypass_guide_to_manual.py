from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_offset_diem_hut.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_bo_sung_setsensor.docx"
image = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-4a594898-8c7c-4150-a7fc-758b3b9f97c6.png")

doc = Document(source)
anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)

created = []
heading = doc.add_paragraph("7.8 Chế độ tạm thời SetSensor", style="Heading 2")
created.append(heading._p)

intro = doc.add_paragraph(
    "SetSensor là chế độ dự phòng tạm thời khi cảm biến xác nhận hút của đầu Tool bị hỏng "
    "và kỹ thuật chưa thể sửa chữa ngay."
)
created.append(intro._p)

picture = doc.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture(str(image), width=Inches(3.0))
picture.paragraph_format.space_after = Pt(3)
created.append(picture._p)

caption = doc.add_paragraph("Hình 10. Tùy chọn SetSensor trong nhóm System")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(7)
for run in caption.runs:
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 101, 112)
created.append(caption._p)

behavior = doc.add_paragraph()
behavior.paragraph_format.space_after = Pt(6)
run = behavior.add_run("KHI BẬT SETSENSOR: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
behavior.add_run(
    "Phần mềm bỏ qua việc xác nhận cảm biến hút của Tool, chờ theo thời gian cài đặt và coi thao tác hút là thành công. "
    "Chế độ này không xác nhận được sản phẩm có thực sự được giữ trên đầu hút hay không."
)
created.append(behavior._p)

scope = doc.add_paragraph()
scope.paragraph_format.space_after = Pt(6)
run = scope.add_run("PHẠM VI: ")
run.bold = True
run.font.color.rgb = RGBColor(22, 74, 138)
scope.add_run(
    "SetSensor chỉ bỏ qua cảm biến xác nhận hút của Tool; không thay thế hoặc vô hiệu hóa kiểm tra cửa an toàn, "
    "áp suất khí tổng, cảm biến Basket hay cảm biến Full Work."
)
created.append(scope._p)

steps = [
    ("1. Xác nhận lỗi: ", "Kỹ thuật/người được phân quyền xác nhận cảm biến hút hỏng nhưng cơ cấu hút, ống khí và đầu Tool vẫn hoạt động."),
    ("2. Đánh giá rủi ro: ", "Dọn vùng máy, giảm tốc độ và bố trí người giám sát liên tục."),
    ("3. Bật tạm thời: ", "Trong Settings > System, đánh dấu SetSensor và ghi nhận thời gian, Job, Tool cùng lý do sử dụng."),
    ("4. Theo dõi sản phẩm: ", "Quan sát từng lần hút/thả; dừng ngay nếu đầu hút không có sản phẩm, sản phẩm rơi hoặc vị trí bất thường."),
    ("5. Khôi phục: ", "Sau khi sửa cảm biến, bỏ dấu SetSensor, kiểm tra tín hiệu thật và chạy thử tốc độ thấp trước khi Auto."),
]
for label, detail in steps:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.38)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)
    paragraph.add_run(detail)
    created.append(paragraph._p)

warning = doc.add_paragraph()
warning.paragraph_format.space_after = Pt(8)
run = warning.add_run("CẢNH BÁO NGHIÊM TRỌNG: ")
run.bold = True
run.font.color.rgb = RGBColor(166, 27, 27)
warning.add_run(
    "Không dùng SetSensor như chế độ vận hành bình thường. Không rời máy khi đang bật. "
    "Nếu không thể giám sát trực tiếp hoặc có nguy cơ rơi sản phẩm, phải dừng máy chờ sửa chữa."
)
created.append(warning._p)

previous = anchor._p.getprevious()
for element in created:
    previous.addnext(element)
    previous = element

doc.core_properties.comments = "Bổ sung hướng dẫn SetSensor tạm thời ngày 11/08/2026"
doc.save(output)
print(output)
