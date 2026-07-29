from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"E:\Nittan\WpfCompanyApp_net58")
OUT = ROOT / "output" / "documents"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Huong_dan_van_hanh_KBOT.docx"
LOGO = ROOT / "WpfCompanyApp" / "Image" / "NittanVietnamLogo.jpg"
SCREEN = Path(r"C:\Users\LinhVo\AppData\Local\Temp\codex-clipboard-39efaf88-fc24-46e3-aa99-a1a17db19248.png")

BLUE = "164A8A"
DARK = "17324D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "B07A00"
RED = "A61B1B"
GREEN = "1B6B35"
GRAY = "5B6570"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

for name, size, color, before, after in [
    ("Title", 28, DARK, 0, 8),
    ("Subtitle", 13, GRAY, 0, 10),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 9, 4),
]:
    s = styles[name]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(sum(widths)))
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "120")
    tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(widths[i]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)
    return p

def add_step(number, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.38)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(f"{number}. {title}: ")
    set_run(r, bold=True, color=DARK)
    p.add_run(detail)

def callout(label, text, kind="note"):
    color = {"warning": RED, "caution": GOLD, "note": BLUE, "ok": GREEN}[kind]
    table = doc.add_table(rows=1, cols=1)
    fixed_table(table, [9360])
    c = table.cell(0, 0)
    shade(c, PALE if kind != "warning" else "FBECEC")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + ": ")
    set_run(r, bold=True, color=color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, LIGHT)
        for r in c.paragraphs[0].runs:
            set_run(r, bold=True, color=DARK, size=9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run(r, size=9.3)
    fixed_table(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def page_break():
    doc.add_page_break()

# Running header/footer
header = sec.header
hp = header.paragraphs[0]
hp.text = "KBOT  |  HƯỚNG DẪN VẬN HÀNH"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    set_run(r, size=8.5, color=GRAY, bold=True)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("NITTAN VIETNAM  •  Tài liệu nội bộ  |  Trang ")
set_run(r, size=8, color=GRAY)
add_field(fp, "PAGE")

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(28)
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(2.2))
doc.add_paragraph().paragraph_format.space_after = Pt(26)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HƯỚNG DẪN VẬN HÀNH")
set_run(r, size=28, bold=True, color=DARK)
p.paragraph_format.space_after = Pt(5)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HỆ THỐNG ROBOT GẮP VÀ ĐẶT TỰ ĐỘNG KBOT")
set_run(r, size=17, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(22)
callout("Phạm vi", "Dành cho người vận hành, kỹ thuật viên và nhân viên bảo trì đã được đào tạo.", "note")
doc.add_paragraph().paragraph_format.space_after = Pt(45)
meta = add_table(
    ["Thông tin", "Nội dung"],
    [
        ["Tên phần mềm", "KBOT - Automatic Pick & Place System"],
        ["Đơn vị", "Nittan Vietnam"],
        ["Phiên bản tài liệu", "01"],
        ["Ngày phát hành", "29/07/2026"],
        ["Tình trạng", "Bản vận hành nội bộ"],
    ],
    [2700, 6660],
)
page_break()

# Document control + TOC
doc.add_heading("Kiểm soát tài liệu", level=1)
add_table(
    ["Phiên bản", "Ngày", "Nội dung cập nhật", "Người duyệt"],
    [["01", "29/07/2026", "Phát hành hướng dẫn vận hành KBOT", "........................"]],
    [1300, 1600, 4460, 2000],
)
doc.add_heading("Mục lục", level=1)
toc = [
    ("1", "Quy định an toàn"),
    ("2", "Tổng quan hệ thống và giao diện"),
    ("3", "Kiểm tra trước khi khởi động"),
    ("4", "Vận hành tự động"),
    ("5", "Ý nghĩa các nút điều khiển"),
    ("6", "Vận hành Manual"),
    ("7", "Cài đặt Job và vị trí robot"),
    ("8", "Xử lý cảnh báo và sự cố"),
    ("9", "Dừng máy và tắt hệ thống"),
    ("10", "Kiểm tra và bảo dưỡng hằng ngày"),
    ("Phụ lục", "Phiếu kiểm tra nhanh tại máy"),
]
for n, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{n}. ")
    set_run(r, bold=True, color=BLUE)
    p.add_run(title)
callout("Lưu ý", "Tài liệu này mô tả phần mềm đang triển khai. Quy trình an toàn tại nhà máy và hướng dẫn của nhà sản xuất robot luôn có mức ưu tiên cao hơn.", "caution")
page_break()

# 1 Safety
doc.add_heading("1. Quy định an toàn", level=1)
callout("Cảnh báo", "Robot có thể chuyển động nhanh và bất ngờ. Không vào vùng làm việc khi robot đang được cấp nguồn hoặc chưa được khóa nguồn an toàn.", "warning")
doc.add_heading("1.1 Yêu cầu đối với người vận hành", level=2)
for x in [
    "Chỉ người đã được đào tạo mới được vận hành Auto, Manual, cài đặt Job hoặc dạy điểm robot.",
    "Luôn biết vị trí nút Emergency Stop vật lý trước khi khởi động.",
    "Không vô hiệu hóa cảm biến cửa, cảm biến áp suất khí hoặc mạch an toàn.",
    "Không dùng nút Stop/Pause trên màn hình thay cho Emergency Stop khi có nguy hiểm tức thời.",
    "Không đứng trong vùng chuyển động của robot khi chạy thử điểm hoặc Jog.",
]:
    add_bullet(x)
doc.add_heading("1.2 Khi nào phải nhấn Emergency Stop", level=2)
for x in [
    "Có người hoặc vật lạ đi vào vùng nguy hiểm.",
    "Robot chuyển động sai hướng, va chạm hoặc có nguy cơ va chạm.",
    "Rơi sản phẩm, rò khí lớn, tiếng động bất thường, khói hoặc mùi cháy.",
    "Cửa an toàn/mạch liên động không hoạt động đúng.",
]:
    add_bullet(x)
callout("Sau Emergency Stop", "Xác định và loại bỏ nguyên nhân, đưa người ra khỏi vùng nguy hiểm, nhả Emergency Stop theo quy trình nhà máy, sau đó mới Reset. Không Reset liên tục khi chưa rõ nguyên nhân.", "warning")

# 2 Overview
doc.add_heading("2. Tổng quan hệ thống và giao diện", level=1)
doc.add_paragraph("KBOT điều khiển và giám sát chu trình robot gắp - đặt, Job sản xuất, camera, cơ cấu khí, cảm biến và trạng thái máy.")
if SCREEN.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCREEN), width=Inches(4.2))
    cap = doc.add_paragraph("Hình 1. Nhóm nút vận hành chính trên màn hình Home")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
doc.add_heading("2.1 Thanh điều hướng bên trái", level=2)
add_table(
    ["Mục", "Chức năng", "Đối tượng sử dụng"],
    [
        ["Home", "Theo dõi trạng thái, chọn Job và điều khiển chu trình Auto.", "Người vận hành"],
        ["Settings", "Quản lý Job, vị trí, quỹ đạo, tốc độ và cảm biến.", "Kỹ thuật viên được phân quyền"],
        ["Manual", "Jog robot, điều khiển I/O và kiểm tra trạng thái cảm biến.", "Kỹ thuật viên/người được đào tạo"],
    ],
    [1800, 4860, 2700],
)
doc.add_heading("2.2 Trạng thái máy", level=2)
add_table(
    ["Trạng thái", "Ý nghĩa", "Hành động điển hình"],
    [
        ["Idle / Stop", "Máy đang chờ, chưa chạy chu trình.", "Chọn Job, Home, Start hoặc Reset."],
        ["Running", "Chu trình tự động đang thực hiện.", "Theo dõi; dùng Pause/Stop khi cần."],
        ["Paused", "Chu trình tạm dừng tại trạng thái hiện tại.", "Khôi phục điều kiện an toàn rồi nhấn Start để tiếp tục."],
        ["Error", "Hệ thống phát hiện lỗi và dừng.", "Đọc log, xử lý nguyên nhân, sau đó Reset."],
    ],
    [1900, 3820, 3740],
)

# 3 Prestart
doc.add_heading("3. Kiểm tra trước khi khởi động", level=1)
add_step(1, "Kiểm tra vùng máy", "Không có người, dụng cụ hoặc vật lạ trong vùng chuyển động.")
add_step(2, "Kiểm tra cơ khí", "Đồ gá, khay, đầu hút, xi lanh và ống khí ở trạng thái bình thường.")
add_step(3, "Kiểm tra an toàn", "Cửa bảo vệ đóng; Emergency Stop đã nhả; mạch an toàn sẵn sàng.")
add_step(4, "Kiểm tra khí nén", "Áp suất khí đạt yêu cầu nội bộ và không có rò rỉ bất thường.")
add_step(5, "Kiểm tra kết nối", "Robot, camera, cảm biến và giao tiếp I/O không báo lỗi.")
add_step(6, "Kiểm tra Job", "Tên Job hiển thị trên Home đúng với sản phẩm sẽ chạy.")
add_step(7, "Kiểm tra Home", "Nếu robot chưa ở vị trí an toàn/Home, thực hiện Home theo điều kiện cho phép.")
callout("Điều kiện Start", "Phần mềm kiểm tra các liên động trước khi chạy. Nếu Start bị từ chối, đọc thông báo/log, khôi phục đúng điều kiện và nhấn Start lại; không Reset nếu hệ thống chỉ báo thiếu liên động.", "note")

# 4 Auto
doc.add_heading("4. Vận hành tự động", level=1)
doc.add_heading("4.1 Khởi động chu trình", level=2)
add_step(1, "Mở màn hình Home", "Nhấn biểu tượng Home trên thanh điều hướng.")
add_step(2, "Chọn Job", "Nhấn “Chọn Job”, chọn đúng công thức sản phẩm và xác nhận tên Job hiển thị.")
add_step(3, "Đưa robot về Home", "Khi máy ở Idle và vùng máy an toàn, nhấn Home nếu cần.")
add_step(4, "Xác nhận sẵn sàng", "Kiểm tra cửa, khí nén, robot, camera, khay và phôi.")
add_step(5, "Bắt đầu", "Nhấn Start một lần. Quan sát chu trình đầu tiên và log trạng thái.")
doc.add_heading("4.2 Tạm dừng và tiếp tục", level=2)
add_bullet("Nhấn Pause để tạm dừng chu trình tại trạng thái hiện tại.")
add_bullet("Không vào vùng máy chỉ vì màn hình hiển thị Paused; phải áp dụng khóa an toàn theo quy định.")
add_bullet("Sau khi khôi phục cửa và các liên động, nhấn Start để tiếp tục.")
doc.add_heading("4.3 Dừng chu trình", level=2)
add_bullet("Nhấn Stop khi cần dừng chu trình có kiểm soát.")
add_bullet("Quan sát trạng thái máy trở về Idle/Stop trước khi thao tác tiếp.")
add_bullet("Nếu có nguy hiểm tức thời, dùng Emergency Stop vật lý, không chờ Stop phần mềm.")

# 5 buttons
doc.add_heading("5. Ý nghĩa các nút điều khiển", level=1)
add_table(
    ["Nút", "Chức năng", "Điều kiện/lưu ý"],
    [
        ["Start", "Bắt đầu chu trình Auto hoặc tiếp tục từ Paused.", "Chỉ chạy khi các liên động và dữ liệu Job hợp lệ."],
        ["Stop", "Yêu cầu dừng chu trình.", "Chờ trạng thái Idle trước thao tác tiếp theo."],
        ["Home", "Đưa robot về vị trí Home theo chương trình.", "Chỉ dùng khi vùng máy trống và trạng thái cho phép."],
        ["Pause", "Tạm dừng chu trình hiện tại.", "Start để tiếp tục sau khi điều kiện an toàn được phục hồi."],
        ["Reset", "Xóa lỗi sau khi nguyên nhân đã được xử lý.", "Không có tác dụng khi máy đang Running; không nhấn lặp lại để bỏ qua lỗi."],
        ["Restart", "Khởi động lại ứng dụng/hệ thống theo xác nhận.", "Chỉ thực hiện khi máy Idle."],
        ["Shutdown", "Tắt máy tính điều khiển theo xác nhận.", "Chỉ thực hiện khi máy Idle và đã kết thúc sản xuất."],
        ["Camera", "Mở/kiểm tra chức năng camera.", "Dùng theo quy trình kiểm tra camera."],
        ["Clear", "Xóa nội dung log hiển thị.", "Không sửa nguyên nhân lỗi; ghi lại thông tin trước khi xóa nếu cần."],
    ],
    [1500, 3900, 3960],
)

# 6 Manual
doc.add_heading("6. Vận hành Manual", level=1)
callout("Cảnh báo", "Manual cho phép điều khiển trực tiếp robot và cơ cấu chấp hành. Chỉ sử dụng ở tốc độ an toàn, từng thao tác một, với vùng máy được kiểm soát.", "warning")
doc.add_heading("6.1 Các nhóm chức năng", level=2)
add_table(
    ["Nhóm", "Chức năng"],
    [
        ["Camera", "Trigger camera để kiểm tra chụp/nhận dạng."],
        ["Lamp / Buzzer", "Bật/tắt đèn xanh, đỏ, vàng và còi."],
        ["Blow Air", "Điều khiển thổi khí CO4-CO6."],
        ["Cylinders", "Điều khiển xi lanh DO0-DO2."],
        ["Suction Cups", "Điều khiển hút chân không DO3-DO5."],
        ["Jog Control", "Di chuyển robot theo trục/hướng; có Step mode."],
        ["Robot", "Power, Free Drive, Reset Robot và đọc Status Robot."],
        ["Sensor IO", "Theo dõi cảm biến xi lanh, hút, cửa, khay, máy và áp suất khí."],
    ],
    [2500, 6860],
)
doc.add_heading("6.2 Quy trình Jog robot", level=2)
add_step(1, "Chuyển sang Manual", "Xác nhận Auto đã dừng và máy ở trạng thái an toàn.")
add_step(2, "Giảm tốc độ", "Chọn tốc độ thấp phù hợp trước lần Jog đầu tiên.")
add_step(3, "Kiểm tra hướng", "Xác định rõ trục/hướng sẽ di chuyển và khoảng trống xung quanh.")
add_step(4, "Jog ngắn", "Nhấn từng nhịp ngắn hoặc dùng Step mode; luôn quan sát robot.")
add_step(5, "Dừng ngay", "Nhả nút khi chuyển động không đúng; dùng Emergency Stop nếu có nguy hiểm.")
callout("Free Drive", "Chỉ bật khi robot được đỡ an toàn, không mang tải nguy hiểm và người thao tác hiểu rõ cơ chế trợ lực của robot.", "caution")

# 7 Settings
doc.add_heading("7. Cài đặt Job và vị trí robot", level=1)
callout("Phân quyền", "Settings thay đổi dữ liệu sản xuất và quỹ đạo. Người vận hành thông thường chỉ chọn Job đã được kỹ thuật phê duyệt.", "caution")
doc.add_heading("7.1 Quản lý Job", level=2)
for x in [
    "Add Job: tạo Job mới.",
    "Reload: nạp lại danh sách Job từ dữ liệu.",
    "Save Job: lưu thông tin Job hiện tại.",
    "Delete: xóa Job; phải xác nhận đúng Job và có bản sao dữ liệu cần thiết.",
    "Export TXT: xuất dữ liệu để kiểm tra/lưu trữ.",
]:
    add_bullet(x)
doc.add_heading("7.2 Dạy điểm và quỹ đạo", level=2)
for x in [
    "Save Home / Save Pick / Save P1...P10: ghi vị trí robot hiện tại vào điểm tương ứng.",
    "Move Home / Move Pick / Move P1...P10: chạy thử tới điểm đã lưu.",
    "Đi Thả 1...5 và Quay Về 1...5: lưu các điểm quỹ đạo đi và về.",
    "Save Vel: lưu tốc độ cho đoạn chuyển động tương ứng.",
]:
    add_bullet(x)
callout("Nguyên tắc chạy thử", "Sau khi thay đổi điểm, chạy thử ở Manual với tốc độ thấp, không có sản phẩm nếu có thể; xác nhận từng điểm và khoảng hở trước khi cho Auto.", "warning")
doc.add_heading("7.3 Cảm biến và lựa chọn công nghệ", level=2)
doc.add_paragraph("Màn hình SETTING ROBOT SENSOR cho phép chọn Basket, Tool 1-3, tốc độ robot, thổi khí và các tùy chọn hệ thống. Chỉ thay đổi theo phiếu thông số đã được phê duyệt.")

# 8 Troubleshooting
doc.add_heading("8. Xử lý cảnh báo và sự cố", level=1)
doc.add_heading("8.1 Trình tự chung", level=2)
add_step(1, "Đảm bảo an toàn", "Dừng máy; dùng Emergency Stop nếu có nguy hiểm.")
add_step(2, "Đọc thông báo", "Ghi lại trạng thái, nội dung log, Job và bước chu trình.")
add_step(3, "Kiểm tra nguyên nhân", "Cửa, khí nén, cảm biến, phôi, đầu hút, robot, camera và kết nối.")
add_step(4, "Khắc phục", "Loại bỏ nguyên nhân theo quy trình kỹ thuật.")
add_step(5, "Reset", "Chỉ nhấn Reset sau khi nguyên nhân đã được xử lý.")
add_step(6, "Chạy xác nhận", "Home nếu cần, chạy thử chu trình đầu và quan sát.")
add_table(
    ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý an toàn"],
    [
        ["Start không chạy", "Cửa mở, thiếu khí, robot/camera chưa sẵn sàng, Job hoặc quỹ đạo thiếu.", "Đọc log, khôi phục liên động, xác nhận Job rồi Start lại."],
        ["Máy chuyển Paused", "Người vận hành nhấn Pause hoặc cửa/liên động bị mất khi chạy.", "Không vào vùng máy; khôi phục điều kiện rồi Start để tiếp tục."],
        ["Robot báo Error", "Lỗi controller, va chạm, Emergency Stop hoặc lỗi chuyển động.", "Dừng, kiểm tra robot và vùng máy; xử lý nguyên nhân rồi Reset Robot/Reset theo quy trình."],
        ["Không hút được sản phẩm", "Thiếu chân không, ống hút rò, giác hút lệch, cảm biến hút không xác nhận.", "Dừng, kiểm tra đầu hút/ống/sản phẩm và cảm biến; không tăng thời gian hút tùy ý."],
        ["Xi lanh không xác nhận", "Thiếu khí, kẹt cơ khí, cảm biến hành trình hoặc dây tín hiệu.", "Ngắt chuyển động, kiểm tra cơ khí/khí/cảm biến; thử Manual sau khi an toàn."],
        ["Camera không phản hồi", "Mất kết nối, trigger/đèn hoặc chương trình camera.", "Kiểm tra kết nối và Trigger trong Manual/Settings; gọi kỹ thuật nếu lặp lại."],
        ["Reset không có tác dụng", "Máy còn Running hoặc nguyên nhân lỗi chưa hết.", "Stop/Pause theo tình huống, xử lý nguyên nhân và chờ Idle/Error rồi Reset."],
    ],
    [2200, 3330, 3830],
)
callout("Không tự xử lý", "Dừng và báo kỹ thuật khi lỗi lặp lại, có va chạm, sai quỹ đạo, mất chức năng an toàn, rò khí lớn, mùi khét hoặc thiết bị quá nhiệt.", "warning")

# 9 shutdown
doc.add_heading("9. Dừng máy và tắt hệ thống", level=1)
add_step(1, "Kết thúc chu trình", "Chờ hoàn tất sản phẩm hiện tại hoặc nhấn Stop để dừng có kiểm soát.")
add_step(2, "Xác nhận Idle", "Kiểm tra trạng thái máy đã về Idle/Stop.")
add_step(3, "Đưa về Home", "Nếu quy trình ca yêu cầu và vùng máy an toàn, nhấn Home.")
add_step(4, "Ghi nhận sản xuất", "Lưu số liệu, lỗi và các bất thường của ca.")
add_step(5, "Tắt phần mềm/máy tính", "Nhấn Shutdown và xác nhận; không ngắt nguồn đột ngột.")
add_step(6, "Cô lập năng lượng", "Tắt nguồn/khí theo quy trình nhà máy khi bảo trì hoặc hết ca.")
callout("Restart", "Chỉ dùng Restart khi cần khởi động lại ứng dụng/hệ điều hành và máy đang Idle. Sau khi khởi động lại phải thực hiện lại kiểm tra trước vận hành.", "note")

# 10 maintenance
doc.add_heading("10. Kiểm tra và bảo dưỡng hằng ngày", level=1)
add_table(
    ["Thời điểm", "Hạng mục", "Tiêu chí"],
    [
        ["Đầu ca", "Vùng máy và che chắn", "Sạch, không vật lạ, cửa và khóa an toàn nguyên vẹn."],
        ["Đầu ca", "Khí nén/ống khí", "Đủ áp, không rò, không gập/nứt."],
        ["Đầu ca", "Đầu hút và xi lanh", "Sạch, không mòn/nứt, chuyển động trơn."],
        ["Đầu ca", "Camera/đèn", "Ống kính sạch, trigger và chiếu sáng ổn định."],
        ["Trong ca", "Log và trạng thái", "Không có lỗi lặp lại hoặc thời gian chu trình bất thường."],
        ["Cuối ca", "Khu vực robot", "Thu dọn sản phẩm, phế phẩm và dụng cụ."],
        ["Cuối ca", "Bàn giao", "Ghi Job, sản lượng, lỗi và xử lý đã thực hiện."],
    ],
    [1700, 3100, 4560],
)
doc.add_heading("Thông tin cần cung cấp khi báo kỹ thuật", level=2)
for x in [
    "Thời gian xảy ra lỗi và tên Job đang chạy.",
    "Trạng thái máy: Idle, Running, Paused hay Error.",
    "Nguyên văn thông báo/log gần thời điểm lỗi.",
    "Vị trí robot, sản phẩm và cơ cấu đang hoạt động.",
    "Hình ảnh/video hiện tượng nếu việc ghi hình an toàn.",
    "Các bước đã thực hiện trước khi lỗi xuất hiện.",
]:
    add_bullet(x)

# Appendix quick checklist
doc.add_heading("Phụ lục A. Phiếu kiểm tra nhanh tại máy", level=1)
add_table(
    ["STT", "Nội dung xác nhận", "Đạt", "Không đạt"],
    [
        ["1", "Vùng robot không có người/vật lạ", "☐", "☐"],
        ["2", "Cửa và mạch an toàn hoạt động", "☐", "☐"],
        ["3", "Emergency Stop đã nhả và sẵn sàng", "☐", "☐"],
        ["4", "Áp suất khí và hệ thống hút bình thường", "☐", "☐"],
        ["5", "Robot, camera và I/O không báo lỗi", "☐", "☐"],
        ["6", "Đúng Job và đúng sản phẩm", "☐", "☐"],
        ["7", "Robot ở vị trí an toàn/Home", "☐", "☐"],
        ["8", "Đã quan sát chu trình đầu tiên", "☐", "☐"],
    ],
    [800, 6160, 1200, 1200],
)
doc.add_paragraph("Ngày: ____/____/________     Ca: ______     Job: ______________________________")
doc.add_paragraph("Người vận hành: __________________________     Ký xác nhận: __________________________")
doc.add_paragraph("Ghi chú bất thường:")
for _ in range(3):
    p = doc.add_paragraph("_" * 104)
    p.paragraph_format.space_after = Pt(3)
callout("Số liên hệ kỹ thuật", "Điền số nội bộ và tên người phụ trách tại đây trước khi phát hành bản chính thức.", "note")

# Keep table rows together where possible and prevent heading orphans.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                p.paragraph_format.widow_control = True

# Metadata
doc.core_properties.title = "Hướng dẫn vận hành hệ thống KBOT"
doc.core_properties.subject = "Vận hành phần mềm và máy robot gắp đặt tự động"
doc.core_properties.author = "Nittan Vietnam"
doc.core_properties.keywords = "KBOT, vận hành, robot, Nittan Vietnam"
doc.core_properties.comments = "Tài liệu vận hành nội bộ"

doc.save(DOCX)
print(DOCX)
