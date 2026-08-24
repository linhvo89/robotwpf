from pathlib import Path
from docx import Document

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_giai_thich_nut_robot.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_khong_trinh_bay_free_drive.docx"

doc = Document(source)
target = "FREE DRIVE: Chỉ bật khi robot được đỡ an toàn"

# Dòng này nằm trong một ô chú ý một cột; xóa cả khung để không để lại khoảng trống.
for table in list(doc.tables):
    if target in "\n".join(cell.text for row in table.rows for cell in row.cells):
        table._element.getparent().remove(table._element)

# Dự phòng nếu nội dung xuất hiện dưới dạng đoạn văn thường.
for paragraph in list(doc.paragraphs):
    if target in paragraph.text:
        paragraph._element.getparent().remove(paragraph._element)

doc.core_properties.comments = "Đã xóa dòng FREE DRIVE theo yêu cầu ngày 11/08/2026"
doc.save(output)
print(output)
