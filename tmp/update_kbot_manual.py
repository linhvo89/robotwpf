from pathlib import Path
from docx import Document

ROOT = Path(r"E:\Nittan\WpfCompanyApp_net58")
DOCX = ROOT / "output" / "documents" / "Huong_dan_van_hanh_KBOT.docx"

doc = Document(DOCX)

replacements = {
    "KBOT - Automatic Pick & Place System": "KBOT 2.1.1 - Automatic Pick & Place System",
    "29/07/2026": "11/08/2026",
    "Phát hành hướng dẫn vận hành KBOT": "Cập nhật hướng dẫn vận hành cho phần mềm KBOT 2.1.1",
}

def replace_in_paragraph(paragraph):
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated != original:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = updated
        else:
            paragraph.add_run(updated)

for paragraph in doc.paragraphs:
    replace_in_paragraph(paragraph)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph)

# Cập nhật số phiên bản tài liệu ở bảng thông tin bìa và bảng kiểm soát.
for table in doc.tables:
    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells]
        if len(values) >= 2 and values[0] == "Phiên bản tài liệu":
            row.cells[1].text = "02"
        if len(values) >= 4 and values[0] == "01" and "Cập nhật hướng dẫn" in values[2]:
            row.cells[0].text = "02"

doc.core_properties.title = "Hướng dẫn vận hành hệ thống KBOT"
doc.core_properties.subject = "Tài liệu vận hành phần mềm KBOT 2.1.1"
doc.core_properties.comments = "Cập nhật ngày 11/08/2026"
doc.save(DOCX)
print(DOCX)
