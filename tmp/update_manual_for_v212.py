from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor

root = Path(r"E:\Nittan\WpfCompanyApp_net58")
source = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_khong_trinh_bay_free_drive.docx"
output = root / "output" / "documents" / "Huong_dan_van_hanh_KBOT_v2.1.2.docx"

doc = Document(source)

def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        return False
    updated = paragraph.text.replace(old, new)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)
    return True

# Cập nhật phiên bản phần mềm và nội dung kiểm soát tài liệu.
for paragraph in doc.paragraphs:
    replace_text(paragraph, "KBOT 2.1.1", "KBOT 2.1.2")
for table in doc.tables:
    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text(paragraph, "KBOT 2.1.1", "KBOT 2.1.2")
                replace_text(
                    paragraph,
                    "Cập nhật hướng dẫn vận hành cho phần mềm KBOT 2.1.2",
                    "Cập nhật giới hạn và giao diện vận hành cho phần mềm KBOT 2.1.2",
                )
        if len(values) >= 2 and values[0] == "Phiên bản tài liệu":
            row.cells[1].text = "03"
        if len(values) >= 4 and values[0] == "02":
            row.cells[0].text = "03"

# H1-H3: giới hạn mới của chương trình là -10..40 mm.
for paragraph in doc.paragraphs:
    replace_text(
        paragraph,
        "Chỉ cài đặt trong giới hạn từ 0 đến 20.",
        "Chỉ cài đặt trong giới hạn từ −10 đến 40 mm. Phần mềm từ chối giá trị nằm ngoài dải này.",
    )

# Offset: bổ sung giới hạn -10..10 theo chương trình.
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("Delta X:"):
        replace_text(
            paragraph,
            "Số dương dịch theo chiều +X; số âm dịch theo chiều −X.",
            "Số dương dịch theo chiều +X; số âm dịch theo chiều −X. Giới hạn nhập: −10 đến 10.",
        )
    if paragraph.text.startswith("Delta Y:"):
        replace_text(
            paragraph,
            "Số dương dịch theo chiều +Y; số âm dịch theo chiều −Y.",
            "Số dương dịch theo chiều +Y; số âm dịch theo chiều −Y. Giới hạn nhập: −10 đến 10.",
        )

# Free Drive đã bị loại khỏi giao diện v2.1.2: cập nhật bảng chức năng Manual.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text(
                    paragraph,
                    "Power, Free Drive, Reset Robot và đọc Status Robot.",
                    "Power, Reset Robot và đọc Status Robot.",
                )

# Xóa hình giao diện Robot cũ có nút Free Drive và đánh lại số hình phía sau.
robot_caption = next(
    (p for p in doc.paragraphs if p.text.strip() == "Hình 6. Nhóm trạng thái và nút điều khiển Robot trên màn hình Manual"),
    None,
)
if robot_caption is not None:
    previous = robot_caption._p.getprevious()
    if previous is not None:
        previous.getparent().remove(previous)
    robot_caption._p.getparent().remove(robot_caption._p)

figure_updates = {
    "Hình 11. Tùy chọn SetSensor trong nhóm System": "Hình 10. Tùy chọn SetSensor trong nhóm System",
    "Hình 10. Bảng Offset điểm hút cho Basket 1, Basket 2 và Tool 1–3": "Hình 9. Bảng Offset điểm hút cho Basket 1, Basket 2 và Tool 1–3",
    "Hình 9. Chọn cảm biến Full Work của máy nhận sản phẩm": "Hình 8. Chọn cảm biến Full Work của máy nhận sản phẩm",
    "Hình 8. Chọn Basket và các đầu Tool tham gia chu trình": "Hình 7. Chọn Basket và các đầu Tool tham gia chu trình",
    "Hình 7. Nhóm cài đặt tốc độ robot trong màn hình Settings": "Hình 6. Nhóm cài đặt tốc độ robot trong màn hình Settings",
}
for paragraph in doc.paragraphs:
    if paragraph.text.strip() in figure_updates:
        replace_text(paragraph, paragraph.text.strip(), figure_updates[paragraph.text.strip()])

# Thêm giới hạn Jog mới trước mục điều khiển robot.
jog_anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "6.3 Các nút điều khiển và trạng thái Robot"
)
jog_note = doc.add_paragraph()
jog_note.paragraph_format.space_after = Pt(8)
run = jog_note.add_run("GIỚI HẠN STEP MODE V2.1.2: ")
run.bold = True
run.font.color.rgb = RGBColor(176, 122, 0)
jog_note.add_run(
    "Bước dịch chuyển tuyến tính được giới hạn từ 0 đến 20 mm; bước xoay được giới hạn từ 0 đến 2°. "
    "Giá trị vượt giới hạn sẽ tự được đưa về giới hạn gần nhất."
)
jog_anchor._p.addprevious(jog_note._p)

# Thêm ghi chú phân quyền Settings trước mục quản lý Job.
settings_anchor = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "7.1 Quản lý Job"
)
access_note = doc.add_paragraph()
access_note.paragraph_format.space_after = Pt(8)
run = access_note.add_run("PHÂN QUYỀN V2.1.2: ")
run.bold = True
run.font.color.rgb = RGBColor(176, 122, 0)
access_note.add_run(
    "Settings mở mặc định tại SETTING ROBOT SENSOR. Hai trang ROBOT POSITION MANAGEMENT và "
    "SETTING ROBOT TRAJECTORY yêu cầu mật khẩu kỹ thuật; người vận hành không tự ý truy cập hoặc thay đổi."
)
settings_anchor._p.addprevious(access_note._p)

# Bổ sung một bảng tra nhanh các giới hạn mới trước Chương 8.
chapter8 = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip() == "8. Xử lý cảnh báo và sự cố"
    and paragraph.style.name == "Heading 1"
)
limit_heading = doc.add_paragraph("7.9 Bảng giới hạn vận hành phiên bản 2.1.2", style="Heading 2")
limit_table = doc.add_table(rows=1, cols=3)
limit_table.style = "Table Grid"
headers = ["Thông số", "Giới hạn", "Phản ứng của phần mềm"]
for index, value in enumerate(headers):
    limit_table.rows[0].cells[index].text = value
limits = [
    ("H1, H2, H3", "−10 đến 40 mm", "Từ chối giá trị ngoài dải."),
    ("Delta X, Delta Y", "−10 đến 10", "Giữ giá trị trong dải cho phép."),
    ("Step tuyến tính", "0 đến 20 mm", "Tự giới hạn về 0 hoặc 20."),
    ("Step xoay", "0 đến 2°", "Tự giới hạn về 0 hoặc 2°."),
    ("Tốc độ robot", "0,05 đến 1,00", "Chọn theo bước 0,05."),
    ("Áp suất khí vận hành", "0,38 đến 0,60 MPa", "Không vận hành ngoài dải."),
]
for values in limits:
    cells = limit_table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value

note = doc.add_paragraph(
    "Sau khi thay đổi H, Offset, Step hoặc tốc độ, luôn chạy thử Manual ở tốc độ thấp và xác nhận khoảng hở trước khi Auto."
)
note.paragraph_format.space_after = Pt(8)

previous = chapter8._p.getprevious()
for element in (limit_heading._p, limit_table._tbl, note._p):
    previous.addnext(element)
    previous = element

doc.core_properties.title = "Hướng dẫn vận hành hệ thống KBOT 2.1.2"
doc.core_properties.subject = "Tài liệu vận hành cập nhật giới hạn chương trình 2.1.2"
doc.core_properties.comments = "Cập nhật giới hạn chương trình ngày 11/08/2026"
doc.save(output)
print(output)
